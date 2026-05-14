from __future__ import annotations
import asyncio
import csv
import io
import json
import logging
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import pdfplumber
from bs4 import BeautifulSoup
from docx import Document

logger = logging.getLogger(__name__)

IMAGE_EXTENSIONS = frozenset({
    "jpg", "jpeg", "png", "gif", "webp", "bmp", "tiff", "tif",
})

_IMAGE_MIME: dict[str, str] = {
    "jpg": "image/jpeg", "jpeg": "image/jpeg",
    "png": "image/png",  "gif": "image/gif",
    "webp": "image/webp", "bmp": "image/png",
    "tiff": "image/jpeg", "tif": "image/jpeg",
}


@dataclass
class RawDocument:
    content: str
    metadata: dict[str, Any]


def parse_file(file_path: str | Path, file_type: str) -> RawDocument:
    """Synchronous parse — vision OCR is NOT applied here.
    Use :func:`parse_file_async` for vision-enhanced parsing.
    """
    path = Path(file_path)
    ext = file_type.lower().lstrip(".")

    parsers = {
        "pdf":  _parse_pdf,
        "docx": _parse_docx,
        "html": _parse_html,
        "htm":  _parse_html,
        "json": _parse_json,
        "csv":  _parse_csv,
    }

    parser = parsers.get(ext)
    if parser:
        return parser(path)

    # Image types — no text content in sync path
    if ext in IMAGE_EXTENSIONS:
        return RawDocument(
            content=f"[图片文件: {path.name}，视觉模型未启用，无法提取文字内容]",
            metadata={"file_type": ext},
        )

    # Fallback: plain text (txt, md, etc.)
    content = path.read_text(encoding="utf-8", errors="replace")
    return RawDocument(content=content, metadata={"file_type": ext})


async def parse_file_async(file_path: str | Path, file_type: str) -> RawDocument:
    """Async parse with optional vision OCR for PDF and DOCX.

    When ``VISION_ENABLED=true`` in .env, image-heavy PDF pages and DOCX
    embedded images are sent to the configured vision LLM for OCR / description.
    Falls back to sync parse for all other file types.
    """
    from agents.vision_agent import is_vision_enabled

    path = Path(file_path)
    ext = file_type.lower().lstrip(".")

    # Images: always route through vision (falls back to placeholder if disabled)
    if ext in IMAGE_EXTENSIONS:
        return await _parse_image_async(path, ext)

    if is_vision_enabled():
        if ext == "pdf":
            return await parse_pdf_with_vision(path)
        if ext == "docx":
            return await parse_docx_with_vision(path)

    # For all other types (or when vision disabled), run sync parser in thread pool
    return await asyncio.get_event_loop().run_in_executor(
        None, parse_file, file_path, file_type
    )


async def _parse_image_async(path: Path, ext: str) -> RawDocument:
    """Describe / OCR an image file via vision LLM."""
    from agents.vision_agent import describe_image, is_vision_enabled

    if not is_vision_enabled():
        return RawDocument(
            content=f"[图片文件: {path.name}，视觉模型未启用，无法提取文字内容]",
            metadata={"file_type": ext},
        )

    mime = _IMAGE_MIME.get(ext, "image/jpeg")
    try:
        image_bytes = path.read_bytes()
        description = await describe_image(image_bytes, mime=mime)  # type: ignore[arg-type]
        return RawDocument(
            content=description,
            metadata={"file_type": ext, "source_type": "image"},
        )
    except Exception as exc:
        logger.warning("Image vision parse failed for %s: %s", path.name, exc)
        return RawDocument(
            content=f"[图片文件: {path.name}，视觉解析失败: {exc}]",
            metadata={"file_type": ext},
        )


def parse_text(content: str) -> RawDocument:
    return RawDocument(content=content, metadata={"file_type": "text"})


def parse_url_content(html: str, url: str) -> RawDocument:
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()
    text = soup.get_text(separator="\n", strip=True)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return RawDocument(content=text, metadata={"file_type": "url", "source_url": url})


def _parse_pdf(path: Path) -> RawDocument:
    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text)
    return RawDocument(
        content="\n\n".join(pages),
        metadata={"file_type": "pdf", "page_count": len(pages)},
    )


async def parse_pdf_with_vision(path: Path) -> RawDocument:
    """PDF parser that falls back to vision OCR for image-heavy pages.

    Requires VISION_ENABLED=true. Pages with extracted text >= threshold
    are kept as-is; sparse pages are rendered and sent to the vision LLM.
    """
    from config import settings
    from agents.vision_agent import ocr_image

    threshold = settings.vision_pdf_ocr_threshold
    pages: list[str] = []
    ocr_pages = 0

    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = (page.extract_text() or "").strip()
            if len(text) >= threshold:
                pages.append(text)
            else:
                # Render page to PNG and OCR
                try:
                    pil_img = page.to_image(resolution=150).original
                    buf = io.BytesIO()
                    pil_img.save(buf, format="PNG")
                    ocr_text = await ocr_image(buf.getvalue(), mime="image/png")
                    if ocr_text.strip():
                        pages.append(f"[OCR page {i+1}]\n{ocr_text}")
                        ocr_pages += 1
                    elif text:
                        pages.append(text)
                except Exception as exc:
                    logger.warning("PDF vision OCR failed page %d: %s", i + 1, exc)
                    if text:
                        pages.append(text)

    return RawDocument(
        content="\n\n".join(pages),
        metadata={"file_type": "pdf", "page_count": len(pages), "ocr_pages": ocr_pages},
    )


def _parse_docx(path: Path) -> RawDocument:
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return RawDocument(
        content="\n\n".join(paragraphs),
        metadata={"file_type": "docx"},
    )


async def parse_docx_with_vision(path: Path) -> RawDocument:
    """DOCX parser that additionally OCRs embedded images via vision LLM."""
    from agents.vision_agent import ocr_image

    _MIME_MAP = {
        ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".png": "image/png", ".gif": "image/gif", ".webp": "image/webp",
        ".bmp": "image/png",  # convert unsupported mime to png label
    }

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text_body = "\n\n".join(paragraphs)

    # Extract images from the docx zip
    image_texts: list[str] = []
    try:
        with zipfile.ZipFile(path) as z:
            media_names = [n for n in z.namelist() if n.startswith("word/media/")]
            for name in media_names:
                ext = Path(name).suffix.lower()
                mime = _MIME_MAP.get(ext)
                if not mime:
                    continue
                try:
                    img_bytes = z.read(name)
                    ocr_text = await ocr_image(img_bytes, mime=mime)  # type: ignore[arg-type]
                    if ocr_text.strip():
                        image_texts.append(f"[图片: {Path(name).name}]\n{ocr_text}")
                except Exception as exc:
                    logger.warning("DOCX vision OCR failed for %s: %s", name, exc)
    except Exception as exc:
        logger.warning("DOCX image extraction failed: %s", exc)

    combined = text_body
    if image_texts:
        combined += "\n\n" + "\n\n".join(image_texts)

    return RawDocument(
        content=combined,
        metadata={"file_type": "docx", "image_count": len(image_texts)},
    )


def _parse_html(path: Path) -> RawDocument:
    html = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    text = soup.get_text(separator="\n", strip=True)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return RawDocument(content=text, metadata={"file_type": "html"})


def _parse_json(path: Path) -> RawDocument:
    data = json.loads(path.read_text(encoding="utf-8"))
    content = json.dumps(data, ensure_ascii=False, indent=2)
    return RawDocument(content=content, metadata={"file_type": "json"})


def _parse_csv(path: Path) -> RawDocument:
    rows: list[str] = []
    with open(path, newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append(" | ".join(f"{k}: {v}" for k, v in row.items()))
    return RawDocument(
        content="\n".join(rows),
        metadata={"file_type": "csv", "row_count": len(rows)},
    )
